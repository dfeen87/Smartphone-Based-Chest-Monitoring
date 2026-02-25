"""
RespiroSync Dashboard Server
=============================
A lightweight Flask web server that exposes a REST API and serves the
dashboard UI for the RespiroSync phase–memory operator.

Endpoints
---------
  GET  /ping              → keepalive / health-check (unauthenticated)
  GET  /                  → dashboard UI (server/static/index.html)
  POST /api/auth/token    → obtain a JWT bearer token
  GET  /api/status        → system status (uptime, version)          [JWT]
  GET  /api/logs          → recent structured log entries (?n=50)    [JWT]
  GET  /api/config        → current operator configuration           [JWT]
  POST /api/config        → update operator configuration            [JWT]
  POST /api/run           → run the phase–memory operator            [JWT]

Authentication
--------------
Protected endpoints (marked [JWT]) require an ``Authorization: Bearer <token>``
header.  Obtain a token by posting ``{"key": "<API_KEY>"}`` to
``/api/auth/token``.  The expected API key is read from the ``API_KEY``
environment variable (defaults to ``"changeme"`` when not set — override in
production).  Tokens are signed with HS256 using a secret read from the
``JWT_SECRET`` environment variable.

All log output goes to stdout so that Render.com captures it in the
service dashboard without any extra configuration.

Scientific basis: PAPER.md §3–4 (phase–memory operator pipeline).
"""

import csv
import io
import os
import smtplib
import sys
import logging
import threading
import time
from collections import deque
from datetime import datetime, timezone, timedelta
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from functools import wraps
from pathlib import Path

import jwt
from flask import Flask, jsonify, request, send_from_directory, send_file

# ── Import the validation pipeline (validation/ lives at the repo root) ────────
sys.path.insert(0, str(Path(__file__).parent.parent))

from validation.pipeline import (  # noqa: E402
    run_pipeline,
    DEFAULT_MEMORY_SAMPLES,
    DEFAULT_ALPHA,
    DEFAULT_BASELINE_SAMP,
    DEFAULT_FS,
)
from validation.physionet_loader import generate_synthetic_resp  # noqa: E402
from validation.multi_record_validation import (  # noqa: E402
    run_multi_record_validation,
    RESULTS_DIR,
)

# ── Structured logging ──────────────────────────────────────────────────────────
# All handlers write to stdout so Render captures the stream automatically.

LOG_BUFFER_SIZE = 500  # keep the last N log entries in memory for /api/logs

_log_buffer: deque = deque(maxlen=LOG_BUFFER_SIZE)


class _BufferHandler(logging.Handler):
    """Append formatted log records to the in-memory ring buffer."""

    def emit(self, record: logging.LogRecord) -> None:
        import datetime
        ts = datetime.datetime.fromtimestamp(record.created).strftime("%Y-%m-%dT%H:%M:%S")
        _log_buffer.append({
            "ts": ts,
            "level": record.levelname,
            "msg": record.getMessage(),
        })


_fmt = logging.Formatter("%(asctime)s [%(levelname)-8s] %(name)s — %(message)s",
                          datefmt="%Y-%m-%dT%H:%M:%S")

_stream_handler = logging.StreamHandler()  # → stdout → Render log tail
_stream_handler.setFormatter(_fmt)

_buffer_handler = _BufferHandler()
_buffer_handler.setFormatter(_fmt)

logging.basicConfig(level=logging.INFO, handlers=[_stream_handler])

logger = logging.getLogger("respirosync")
logger.addHandler(_buffer_handler)
logger.setLevel(logging.INFO)

# ── Flask application ───────────────────────────────────────────────────────────

app = Flask(__name__, static_folder="static")

_start_time = time.time()

# ── Operator configuration (mutable at runtime via /api/config) ────────────────
_config: dict = {
    "memory_samples":   DEFAULT_MEMORY_SAMPLES,  # M  — rolling window (Eq. 4)
    "alpha":            DEFAULT_ALPHA,            # α  — sensitivity     (Eq. 6)
    "baseline_samples": DEFAULT_BASELINE_SAMP,   # calibration window
    "fs":               DEFAULT_FS,              # sample rate (Hz)
}

# ── Processing state ────────────────────────────────────────────────────────────
_run_lock = threading.Lock()
_last_metrics: dict = {}
_validate_lock = threading.Lock()

# ── JWT configuration ────────────────────────────────────────────────────────────
# JWT_SECRET must be set to a strong random value in production.
# API_KEY is the shared secret clients exchange for a bearer token.
_JWT_SECRET: str = os.environ.get("JWT_SECRET", "changeme-jwt-secret")
_API_KEY: str = os.environ.get("API_KEY", "changeme")
_USING_DEFAULT_KEY: bool = "API_KEY" not in os.environ
_JWT_ALGORITHM = "HS256"
_JWT_EXPIRY_HOURS = 24


def _make_token() -> str:
    """Return a signed JWT valid for *_JWT_EXPIRY_HOURS* hours."""
    payload = {
        "sub": "respirosync",
        "iat": datetime.now(tz=timezone.utc),
        "exp": datetime.now(tz=timezone.utc) + timedelta(hours=_JWT_EXPIRY_HOURS),
    }
    return jwt.encode(payload, _JWT_SECRET, algorithm=_JWT_ALGORITHM)


def require_jwt(f):
    """Decorator: enforce a valid JWT ``Authorization: Bearer <token>`` header."""
    @wraps(f)
    def decorated(*args, **kwargs):
        auth = request.headers.get("Authorization", "")
        if not auth.startswith("Bearer "):
            return jsonify({"error": "Missing or invalid Authorization header"}), 401
        token = auth[len("Bearer "):]
        try:
            jwt.decode(token, _JWT_SECRET, algorithms=[_JWT_ALGORITHM])
        except jwt.ExpiredSignatureError:
            return jsonify({"error": "Token has expired"}), 401
        except jwt.InvalidTokenError:
            return jsonify({"error": "Invalid token"}), 401
        return f(*args, **kwargs)
    return decorated


# ── Routes ──────────────────────────────────────────────────────────────────────

@app.route("/ping")
def ping() -> object:
    """Unauthenticated keepalive endpoint for load-balancers and uptime monitors.

    Also returns ``default_key_active`` so the dashboard can auto-connect when
    no custom ``API_KEY`` environment variable has been configured.
    """
    return jsonify({"pong": True, "default_key_active": _USING_DEFAULT_KEY})


@app.route("/api/auth/token", methods=["POST"])
def api_auth_token() -> object:
    """Exchange the shared API key for a signed JWT bearer token.

    Request body (JSON)
    -------------------
    key : str  The value of the ``API_KEY`` environment variable.

    Response (JSON)
    ---------------
    token : str  A signed JWT to use as ``Authorization: Bearer <token>``.
    """
    body = request.get_json(force=True, silent=True) or {}
    if body.get("key") != _API_KEY:
        logger.warning("Token request rejected — invalid API key")
        return jsonify({"error": "Invalid API key"}), 401
    token = _make_token()
    logger.info("JWT issued for sub=respirosync")
    return jsonify({"token": token})


@app.route("/")
def index() -> object:
    """Serve the dashboard UI."""
    return send_from_directory(app.static_folder, "index.html")


@app.route("/api/status")
@require_jwt
def api_status() -> object:
    """Return current system status."""
    return jsonify({
        "status":    "running",
        "version":   "1.0.0",
        "uptime_s":  round(time.time() - _start_time, 1),
        "pipeline":  "phase-memory operator (PAPER.md §3–4)",
    })


@app.route("/api/logs")
@require_jwt
def api_logs() -> object:
    """Return recent log entries.

    Query params
    ------------
    n : int  Maximum number of entries to return (default 50, max 500).
    """
    try:
        n = min(int(request.args.get("n", 50)), LOG_BUFFER_SIZE)
    except ValueError:
        n = 50
    return jsonify(list(_log_buffer)[-n:])


@app.route("/api/config", methods=["GET"])
@require_jwt
def api_get_config() -> object:
    """Return the current operator configuration."""
    return jsonify(dict(_config))


@app.route("/api/config", methods=["POST"])
@require_jwt
def api_set_config() -> object:
    """Update operator configuration parameters.

    Accepted JSON keys
    ------------------
    memory_samples   : int   (Eq. 4 window M)
    alpha            : float (Eq. 6 sensitivity α ∈ [2, 3])
    baseline_samples : int
    fs               : float (sample rate Hz)
    """
    data = request.get_json(force=True, silent=True) or {}
    updated = {}
    for key, cast in (
        ("memory_samples",   int),
        ("alpha",            float),
        ("baseline_samples", int),
        ("fs",               float),
    ):
        if key in data:
            try:
                _config[key] = cast(data[key])
                updated[key] = _config[key]
            except (ValueError, TypeError) as exc:
                logger.warning("Config update rejected for %s: %s", key, exc)
    if updated:
        logger.info("Configuration updated: %s", updated)
    return jsonify({"status": "ok", "config": dict(_config)})


@app.route("/api/run", methods=["POST"])
@require_jwt
def api_run() -> object:
    """Run the phase–memory operator on a synthetic respiratory signal.

    Returns summary metrics derived from the operator output (PAPER.md §4).
    Use POST body JSON ``{"duration_s": 90}`` to control signal length.
    """
    if not _run_lock.acquire(blocking=False):
        logger.warning("Run requested while another run is in progress — rejected")
        return jsonify({"error": "A run is already in progress"}), 429

    try:
        body = request.get_json(force=True, silent=True) or {}
        duration_s = float(body.get("duration_s", 90))

        logger.info(
            "Operator run started — duration=%.0fs  M=%d  α=%.2f  fs=%.0f Hz",
            duration_s, _config["memory_samples"], _config["alpha"], _config["fs"],
        )

        # Generate synthetic respiratory signal (PAPER.md §5.1 regimes)
        data = generate_synthetic_resp(duration_s=duration_s, fs=int(_config["fs"]))
        signal = data["signal"]
        fs = data["fs"]

        # Run the full phase–memory pipeline (PAPER.md §7.1)
        result = run_pipeline(
            signal,
            fs=float(fs),
            M=int(_config["memory_samples"]),
            alpha=float(_config["alpha"]),
            baseline_samples=int(_config["baseline_samples"]),
        )

        metrics = {
            "sigma_omega":       round(float(result["sigma_omega"]),  6),
            "threshold":         round(float(result["threshold"]),    6),
            "instability_count": int(result["instability"].sum()),
            "instability_rate":  round(float(result["instability"].mean()), 4),
            "delta_phi_max":     round(float(result["delta_phi"].max()),    4),
            "delta_phi_mean":    round(float(result["delta_phi"].mean()),   4),
            "n_samples":         int(len(signal)),
            "duration_s":        round(float(len(signal) / fs), 1),
        }

        _last_metrics.clear()
        _last_metrics.update(metrics)

        logger.info(
            "Operator run complete — σ_ω=%.4f rad/s  threshold=%.4f rad/s  "
            "instability_rate=%.2f%%  n_alarms=%d",
            metrics["sigma_omega"],
            metrics["threshold"],
            metrics["instability_rate"] * 100,
            metrics["instability_count"],
        )

        return jsonify({"status": "ok", "metrics": metrics})

    except Exception as exc:  # pylint: disable=broad-except
        logger.error("Operator run failed: %s", exc, exc_info=True)
        return jsonify({"error": str(exc)}), 500

    finally:
        _run_lock.release()


# ── Validation helpers ──────────────────────────────────────────────────────────

def _read_results_csv(filename: str) -> list:
    """Read a CSV from RESULTS_DIR and return list of row dicts."""
    path = RESULTS_DIR / filename
    if not path.exists():
        return []
    with open(path, newline="") as fh:
        return list(csv.DictReader(fh))


def _build_pdf(rows: list, summary_rows: list, n_records: int) -> bytes:
    """Generate a PDF report from validation results using reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer,
        )
        from reportlab.lib.styles import getSampleStyleSheet
    except ImportError:
        raise RuntimeError("reportlab is required for PDF generation. Install with: pip install reportlab")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("RespiroSync Validation Results", styles["Title"]))
    elements.append(Spacer(1, 12))
    methods = (
        f"Results are averaged across N = {n_records} BIDMC recordings using "
        "the semi-synthetic perturbation protocol described in Section 5."
    )
    elements.append(Paragraph(methods, styles["Normal"]))
    elements.append(Spacer(1, 20))

    # Table 1 — per-record metrics
    elements.append(Paragraph("Table 1: Per-Record Metrics", styles["Heading2"]))
    if rows:
        header = list(rows[0].keys())
        table_data = [header] + [[r.get(k, "") or "—" for k in header] for r in rows]
        t = Table(table_data, hAlign="LEFT")
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4f8ef7")),
            ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
            ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",   (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4ff")]),
            ("GRID",       (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("ALIGN",      (0, 0), (-1, -1), "CENTER"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        elements.append(t)
    else:
        elements.append(Paragraph("No per-record results available.", styles["Normal"]))
    elements.append(Spacer(1, 20))

    # Table 2 — aggregated stats
    elements.append(Paragraph("Table 2: Aggregated Statistics (Mean ± SD)", styles["Heading2"]))
    if summary_rows:
        header2 = list(summary_rows[0].keys())
        table_data2 = [header2] + [[r.get(k, "") or "—" for k in header2] for r in summary_rows]
        t2 = Table(table_data2, hAlign="LEFT")
        t2.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#3ecf8e")),
            ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
            ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",   (0, 0), (-1, -1), 9),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0fff8")]),
            ("GRID",       (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("ALIGN",      (0, 0), (-1, -1), "CENTER"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        elements.append(t2)
    else:
        elements.append(Paragraph("No summary results available.", styles["Normal"]))

    doc.build(elements)
    return buf.getvalue()


def _build_docx(rows: list, summary_rows: list, n_records: int) -> bytes:
    """Generate a DOCX report from validation results using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx")

    doc = Document()
    doc.add_heading("RespiroSync Validation Results", 0)
    methods = (
        f"Results are averaged across N = {n_records} BIDMC recordings using "
        "the semi-synthetic perturbation protocol described in Section 5."
    )
    doc.add_paragraph(methods)
    doc.add_paragraph()

    # Table 1 — per-record metrics
    doc.add_heading("Table 1: Per-Record Metrics", level=2)
    if rows:
        header = list(rows[0].keys())
        tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
        tbl.style = "Light List Accent 1"
        for j, h in enumerate(header):
            tbl.rows[0].cells[j].text = h
            tbl.rows[0].cells[j].paragraphs[0].runs[0].bold = True
        for i, row in enumerate(rows, start=1):
            for j, k in enumerate(header):
                tbl.rows[i].cells[j].text = str(row.get(k, "") or "—")
    else:
        doc.add_paragraph("No per-record results available.")
    doc.add_paragraph()

    # Table 2 — aggregated stats
    doc.add_heading("Table 2: Aggregated Statistics (Mean ± SD)", level=2)
    if summary_rows:
        header2 = list(summary_rows[0].keys())
        tbl2 = doc.add_table(rows=1 + len(summary_rows), cols=len(header2))
        tbl2.style = "Light List Accent 1"
        for j, h in enumerate(header2):
            tbl2.rows[0].cells[j].text = h
            tbl2.rows[0].cells[j].paragraphs[0].runs[0].bold = True
        for i, row in enumerate(summary_rows, start=1):
            for j, k in enumerate(header2):
                tbl2.rows[i].cells[j].text = str(row.get(k, "") or "—")
    else:
        doc.add_paragraph("No summary results available.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Validation routes ───────────────────────────────────────────────────────────

@app.route("/api/validate", methods=["POST"])
@require_jwt
def api_validate() -> object:
    """Run multi-record BIDMC validation (PAPER.md §5).

    Request body (JSON)
    -------------------
    n_records  : int   Number of records to evaluate (default 5, min 1).
    synthetic  : bool  Use synthetic signal fallback (default true).

    Response (JSON)
    ---------------
    stats            : dict  Mean ± SD for each metric.
    n_records        : int   Records processed.
    methods_statement: str   One-sentence Methods text for the paper.
    """
    if not _validate_lock.acquire(blocking=False):
        logger.warning("Validation requested while another validation is in progress — rejected")
        return jsonify({"error": "A validation run is already in progress"}), 429

    try:
        body = request.get_json(force=True, silent=True) or {}
        n_records = max(1, int(body.get("n_records", 5)))
        use_synthetic = bool(body.get("synthetic", True))

        logger.info(
            "Multi-record validation started — n_records=%d  synthetic=%s",
            n_records, use_synthetic,
        )

        result = run_multi_record_validation(
            n_records=n_records,
            use_synthetic=use_synthetic,
        )

        logger.info(
            "Multi-record validation complete — n_records=%d  "
            "drift_latency_mean=%.4f  pause_latency_mean=%.4f",
            result["n_records"],
            result["stats"]["drift_latency"]["mean"] or float("nan"),
            result["stats"]["pause_latency"]["mean"] or float("nan"),
        )

        return jsonify({"status": "ok", **result})

    except Exception as exc:  # pylint: disable=broad-except
        logger.error("Validation run failed: %s", exc, exc_info=True)
        return jsonify({"error": str(exc)}), 500

    finally:
        _validate_lock.release()


@app.route("/api/results/metrics.csv")
@require_jwt
def api_results_metrics_csv() -> object:
    """Download the per-record metrics CSV."""
    path = RESULTS_DIR / "metrics.csv"
    if not path.exists():
        return jsonify({"error": "metrics.csv not found — run /api/validate first"}), 404
    return send_file(str(path), mimetype="text/csv", as_attachment=True,
                     download_name="metrics.csv")


@app.route("/api/results/summary.csv")
@require_jwt
def api_results_summary_csv() -> object:
    """Download the aggregated summary CSV."""
    path = RESULTS_DIR / "summary.csv"
    if not path.exists():
        return jsonify({"error": "summary.csv not found — run /api/validate first"}), 404
    return send_file(str(path), mimetype="text/csv", as_attachment=True,
                     download_name="summary.csv")


@app.route("/api/results/pdf")
@require_jwt
def api_results_pdf() -> object:
    """Download an auto-generated PDF report of the validation results."""
    rows = _read_results_csv("metrics.csv")
    summary_rows = _read_results_csv("summary.csv")
    if not rows and not summary_rows:
        return jsonify({"error": "No results found — run /api/validate first"}), 404
    n_records = len(rows)
    try:
        pdf_bytes = _build_pdf(rows, summary_rows, n_records)
    except RuntimeError as exc:
        return jsonify({"error": str(exc)}), 500
    return send_file(
        io.BytesIO(pdf_bytes), mimetype="application/pdf", as_attachment=True,
        download_name="results.pdf",
    )


@app.route("/api/results/docx")
@require_jwt
def api_results_docx() -> object:
    """Download an auto-generated DOCX report of the validation results."""
    rows = _read_results_csv("metrics.csv")
    summary_rows = _read_results_csv("summary.csv")
    if not rows and not summary_rows:
        return jsonify({"error": "No results found — run /api/validate first"}), 404
    n_records = len(rows)
    try:
        docx_bytes = _build_docx(rows, summary_rows, n_records)
    except RuntimeError as exc:
        return jsonify({"error": str(exc)}), 500
    return send_file(
        io.BytesIO(docx_bytes),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name="results.docx",
    )


@app.route("/api/send-results", methods=["POST"])
@require_jwt
def api_send_results() -> object:
    """Email the validation results (CSV + PDF + DOCX) to the supplied address.

    Request body (JSON)
    -------------------
    email : str  Recipient email address.

    SMTP configuration is read from environment variables:
      SMTP_HOST  (default: localhost)
      SMTP_PORT  (default: 587)
      SMTP_USER  (optional)
      SMTP_PASS  (optional)
      SMTP_FROM  (default: respirosync@localhost)
    """
    body = request.get_json(force=True, silent=True) or {}
    recipient = (body.get("email") or "").strip()
    if not recipient or "@" not in recipient:
        return jsonify({"error": "A valid email address is required"}), 400

    rows = _read_results_csv("metrics.csv")
    summary_rows = _read_results_csv("summary.csv")
    if not rows and not summary_rows:
        return jsonify({"error": "No results found — run /api/validate first"}), 404

    n_records = len(rows)

    # Build attachments
    attachments = []

    # metrics.csv
    metrics_path = RESULTS_DIR / "metrics.csv"
    if metrics_path.exists():
        with open(metrics_path, "rb") as fh:
            attachments.append(("metrics.csv", "text/csv", fh.read()))

    # summary.csv
    summary_path = RESULTS_DIR / "summary.csv"
    if summary_path.exists():
        with open(summary_path, "rb") as fh:
            attachments.append(("summary.csv", "text/csv", fh.read()))

    # results.pdf
    try:
        pdf_bytes = _build_pdf(rows, summary_rows, n_records)
        attachments.append(("results.pdf", "application/pdf", pdf_bytes))
    except Exception as exc:  # pylint: disable=broad-except
        logger.warning("PDF generation skipped for email: %s", exc)

    # results.docx
    try:
        docx_bytes = _build_docx(rows, summary_rows, n_records)
        attachments.append((
            "results.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            docx_bytes,
        ))
    except Exception as exc:  # pylint: disable=broad-except
        logger.warning("DOCX generation skipped for email: %s", exc)

    # Build MIME message
    msg = MIMEMultipart()
    smtp_from = os.environ.get("SMTP_FROM", "respirosync@localhost")
    msg["From"] = smtp_from
    msg["To"] = recipient
    msg["Subject"] = "RespiroSync Validation Results"
    body_text = (
        f"Please find the RespiroSync validation results attached.\n\n"
        f"Results are averaged across N = {n_records} BIDMC recordings using "
        "the semi-synthetic perturbation protocol described in Section 5.\n\n"
        "— RespiroSync"
    )
    msg.attach(MIMEText(body_text, "plain"))

    for filename, mimetype, data in attachments:
        main_type, sub_type = mimetype.split("/", 1)
        part = MIMEBase(main_type, sub_type)
        part.set_payload(data)
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", "attachment", filename=filename)
        msg.attach(part)

    # Send via SMTP
    smtp_host = os.environ.get("SMTP_HOST", "localhost")
    smtp_port = int(os.environ.get("SMTP_PORT", 587))
    smtp_user = os.environ.get("SMTP_USER", "")
    smtp_pass = os.environ.get("SMTP_PASS", "")

    try:
        with smtplib.SMTP(smtp_host, smtp_port, timeout=10) as server:
            server.ehlo()
            if smtp_port != 25:
                server.starttls()
                server.ehlo()
            if smtp_user:
                server.login(smtp_user, smtp_pass)
            server.sendmail(smtp_from, [recipient], msg.as_string())
        logger.info("Validation results emailed to %s", recipient)
        return jsonify({"status": "ok", "sent_to": recipient})
    except Exception as exc:  # pylint: disable=broad-except
        logger.error("Email delivery failed: %s", exc)
        return jsonify({"error": f"Email delivery failed: {exc}"}), 500


# ── Entry point ─────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    logger.info("RespiroSync dashboard starting on port %d", port)
    logger.info("Pipeline: preprocessing → analytic signal → θ(t) → ω(t) → ω̄(t) → ΔΦ(t)")
    app.run(host="0.0.0.0", port=port, debug=False)
